import os
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext
from openpyxl import Workbook, load_workbook
from docx import Document
import smtplib
from email.message import EmailMessage
# pywhatkit imported lazily inside send functions (avoids SSL startup crash)
import pyautogui
import time
import subprocess
from PIL import Image, ImageTk
from datetime import datetime

# --- CONFIGURATION ---
_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_FILE = os.path.join(_BASE_DIR, "Rent_Management.xlsx")
COMPANY_NAME = "Owais Rent Management"
BG_IMAGE_PATH = os.path.join(_BASE_DIR, "rent_bg.png")
LOGO_IMAGE_PATH = os.path.join(_BASE_DIR, "rent_logo.png")

# ── KE Electricity Charges (exact values from bill) ──────────────────────
# Slab rates: 0-100 units = 10.54/unit, 101-200 units = 13.01/unit
KE_SLABS = [
    (100, 10.54),
    (200, 13.01),
    (float('inf'), 13.01),   # beyond 200 — same top rate
]
FIXED_CHARGE     = 600.00    # Fixed monthly charge
ADDITIONAL_CHG   = 169.00   # Additional charges (editable per bill)
ELEC_DUTY        = 19.91    # Electricity duty
SALES_TAX        = 350.53   # Sales tax
KMC_CHG          = 20.00    # KMC charges

def calc_ke_bill(units, additional=ADDITIONAL_CHG):
    """Calculate KE electricity bill with exact slab rates and fixed charges."""
    energy = 0.0
    remaining = units
    prev_limit = 0
    for limit, rate in KE_SLABS:
        if remaining <= 0:
            break
        slab_units = min(remaining, limit - prev_limit)
        energy += slab_units * rate
        remaining -= slab_units
        prev_limit = limit
    total = energy + FIXED_CHARGE + additional + ELEC_DUTY + SALES_TAX + KMC_CHG
    return (
        round(energy, 2),
        round(FIXED_CHARGE, 2),
        round(additional, 2),
        round(ELEC_DUTY, 2),
        round(SALES_TAX, 2),
        round(KMC_CHG, 2),
        round(total, 2)
    )

# --- HELPER FUNCTIONS ---
def sanitize_filename(name: str) -> str:
    """Remove/replace characters invalid in Windows file names."""
    import re
    return re.sub(r'[\\/:*?"<>|]', '_', name).strip() or "Unknown"
def load_bg_img(frame, width=500, height=300):
    try:
        img_frame = tk.Frame(frame, bg="white")
        img_frame.pack(expand=True)
        img_frame.tkraise()
        bg_img = Image.open(BG_IMAGE_PATH)
        bg_img = bg_img.resize((width, height), Image.Resampling.LANCZOS)
        img_tk = ImageTk.PhotoImage(bg_img)
        lbl = tk.Label(img_frame, image=img_tk, bg="white")
        lbl.image = img_tk
        lbl.pack(padx=3, pady=3)
    except Exception as e:
        print(f"Image load error: {e}")
        tk.Label(frame, text="[Image Not Found]", bg="white").pack()

def img(R_frame):
    for widget in R_frame.winfo_children():
        widget.destroy()
    load_bg_img(R_frame)

# --- ELECTRICITY BILL MODULE ---

def electricity_bill_gui(R_frame):
    for widget in R_frame.winfo_children():
        widget.destroy()

    # ---- Styling ----
    SIDEBAR_BG = "#1c2e4a"
    ACCENT     = "#f0a500"
    CARD_BG    = "#f7f9fc"
    TEXT_DARK  = "#1a1a2e"
    BTN_GREEN  = "#28a745"
    BTN_BLUE   = "#007bff"
    BTN_RED    = "#dc3545"
    WA_GREEN   = "#25D366"

    outer = tk.Frame(R_frame, bg=CARD_BG)
    outer.pack(fill="both", expand=True)

    # Header
    hdr = tk.Frame(outer, bg=SIDEBAR_BG, pady=12)
    hdr.pack(fill="x")
    tk.Label(hdr, text="⚡  Electricity Bill Calculator", font=("Segoe UI", 17, "bold"),
             bg=SIDEBAR_BG, fg=ACCENT).pack()
    tk.Label(hdr, text=COMPANY_NAME, font=("Segoe UI", 10),
             bg=SIDEBAR_BG, fg="white").pack()

    # --- Input Area ---
    input_card = tk.Frame(outer, bg="white", relief="groove", bd=1)
    input_card.pack(fill="x", padx=30, pady=15)

    tk.Label(input_card, text="Meter Reading Entry", font=("Segoe UI", 12, "bold"),
             bg="white", fg=TEXT_DARK).grid(row=0, column=0, columnspan=4, pady=(12, 8), padx=20, sticky="w")

    # Row 1 — Customer & Month
    def labeled_entry(parent, r, c, label, width=18, default=""):
        tk.Label(parent, text=label, bg="white", font=("Segoe UI", 10),
                 fg="#555").grid(row=r, column=c, sticky="w", padx=(20,5), pady=6)
        ent = tk.Entry(parent, width=width, font=("Segoe UI", 11),
                       relief="solid", bd=1)
        ent.grid(row=r, column=c+1, sticky="w", padx=(0, 20), pady=6)
        if default:
            ent.insert(0, default)
        return ent

    cust_id_ent  = labeled_entry(input_card, 1, 0, "Customer ID:",
                                 default="LA104056")
    month_ent    = labeled_entry(input_card, 1, 2, "Month/Year:",
                                 default=datetime.now().strftime("%b-%Y"))
    prev_ent     = labeled_entry(input_card, 2, 0, "Previous Reading (kWh):")
    curr_ent     = labeled_entry(input_card, 2, 2, "Current Reading  (kWh):")
    acc_no_ent   = labeled_entry(input_card, 3, 0, "KE Account No:",
                                 default="0400013715589")
    meter_no_ent = labeled_entry(input_card, 3, 2, "Meter Serial No:",
                                 default="SEA78555")
    # Additional Charges editable field (changes every month on KE bill)
    add_chg_ent  = labeled_entry(input_card, 4, 0, "Additional Charges (Rs):",
                                 default=str(ADDITIONAL_CHG))
    cust_phone_ent = labeled_entry(input_card, 4, 2, "Customer WhatsApp No:")

    # --- Result Panel ---
    result_card = tk.Frame(outer, bg="white", relief="groove", bd=1)
    result_card.pack(fill="x", padx=30, pady=5)

    tk.Label(result_card, text="Bill Summary", font=("Segoe UI", 12, "bold"),
             bg="white", fg=TEXT_DARK).grid(row=0, column=0, columnspan=4,
             pady=(12,8), padx=20, sticky="w")

    result_vars = {}
    result_fields = [
        ("Units Consumed:",      "units",      0, 0),
        ("Energy Charges:",      "energy",     0, 2),
        ("Fixed Charge:",        "fixed",      1, 0),
        ("Additional Charges:",  "additional", 1, 2),
        ("Electricity Duty:",    "duty",       2, 0),
        ("Sales Tax:",           "salestax",   2, 2),
        ("KMC Charges:",         "kmc",        3, 0),
        ("TOTAL BILL:",          "total",      3, 2),
    ]
    for label, key, r, c in result_fields:
        tk.Label(result_card, text=label, bg="white", font=("Segoe UI", 10),
                 fg="#555").grid(row=r+1, column=c, sticky="w", padx=(20,5), pady=6)
        var = tk.StringVar(value="---")
        result_vars[key] = var
        bold  = "bold"    if key == "total" else "normal"
        color = BTN_RED   if key == "total" else TEXT_DARK
        tk.Label(result_card, textvariable=var, bg="white",
                 font=("Segoe UI", 11, bold), fg=color).grid(
                 row=r+1, column=c+1, sticky="w", padx=(0,20), pady=6)

    tk.Label(result_card,
             text="Slab: 0-100 units @ Rs.10.54/unit  |  101-200 units @ Rs.13.01/unit",
             bg="white", fg="#888", font=("Segoe UI", 9)).grid(
             row=6, column=0, columnspan=4, pady=(0,12), padx=20, sticky="w")

    # --- Store computed bill ---
    computed = {}

    def calculate():
        try:
            prev = int(prev_ent.get())
            curr = int(curr_ent.get())
        except ValueError:
            messagebox.showerror("Input Error", "Please enter valid numeric meter readings.")
            return
        if curr < prev:
            messagebox.showerror("Input Error", "Current reading cannot be less than previous reading.")
            return
        try:
            add_chg = float(add_chg_ent.get())
        except ValueError:
            add_chg = ADDITIONAL_CHG

        units = curr - prev
        energy, fixed, additional, duty, salestax, kmc, total = calc_ke_bill(units, add_chg)

        result_vars["units"].set(f"{units} kWh")
        result_vars["energy"].set(f"Rs. {energy:,.2f}")
        result_vars["fixed"].set(f"Rs. {fixed:,.2f}")
        result_vars["additional"].set(f"Rs. {additional:,.2f}")
        result_vars["duty"].set(f"Rs. {duty:,.2f}")
        result_vars["salestax"].set(f"Rs. {salestax:,.2f}")
        result_vars["kmc"].set(f"Rs. {kmc:,.2f}")
        result_vars["total"].set(f"Rs. {total:,.2f}")

        computed.update({
            "prev": prev, "curr": curr, "units": units,
            "energy": energy, "fixed": fixed, "additional": additional,
            "duty": duty, "salestax": salestax, "kmc": kmc, "total": total
        })

    def _get_bill_texts(cid, month, acc, meter, cust_name, prop_title):
        """Build plain-text and WhatsApp-formatted bill strings."""
        now = datetime.now().strftime("%d-%b-%Y %H:%M")
        # Plain text for .docx / email preview
        plain = (
            f"\n{'='*55}\n"
            f"         {COMPANY_NAME}\n"
            f"          ELECTRICITY BILL RECEIPT\n"
            f"{'='*55}\n"
            f"  Date Issued : {now}\n"
            f"  Month       : {month}\n"
            f"  Customer ID : {cid}\n"
            f"  Name        : {cust_name}\n"
            f"  Property    : {prop_title}\n"
            f"  KE Acc No   : {acc or 'N/A'}\n"
            f"  Meter No    : {meter or 'N/A'}\n"
            f"{'-'*55}\n"
            f"  Previous Reading : {computed['prev']:>8} kWh\n"
            f"  Current Reading  : {computed['curr']:>8} kWh\n"
            f"  Units Consumed   : {computed['units']:>8} kWh\n"
            f"{'-'*55}\n"
            f"  Energy Charges   : Rs. {computed['energy']:>10,.2f}\n"
            f"  Fixed Charge     : Rs. {computed['fixed']:>10,.2f}\n"
            f"  Additional Chg   : Rs. {computed['additional']:>10,.2f}\n"
            f"  Electricity Duty : Rs. {computed['duty']:>10,.2f}\n"
            f"  Sales Tax        : Rs. {computed['salestax']:>10,.2f}\n"
            f"  KMC Charges      : Rs. {computed['kmc']:>10,.2f}\n"
            f"{'-'*55}\n"
            f"  TOTAL BILL       : Rs. {computed['total']:>10,.2f}\n"
            f"{'='*55}\n"
            f"  Please pay by the due date to avoid disconnection.\n"
            f"  {COMPANY_NAME} — Management\n"
            f"{'='*55}\n"
        )
        # WhatsApp professional format (bold with asterisks)
        wa = (
            f"⚡ *ELECTRICITY BILL — {COMPANY_NAME}*\n"
            f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            f"📅 *Month:* {month}\n"
            f"👤 *Customer:* {cust_name}\n"
            f"🏠 *Property:* {prop_title}\n"
            f"🔢 *KE Account:* {acc or 'N/A'}\n"
            f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            f"📊 *METER READING*\n"
            f"  Previous  : {computed['prev']} kWh\n"
            f"  Current   : {computed['curr']} kWh\n"
            f"  *Units Used : {computed['units']} kWh*\n"
            f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            f"💰 *CHARGES BREAKDOWN*\n"
            f"  Energy Charges   : Rs. {computed['energy']:,.2f}\n"
            f"  Fixed Charge     : Rs. {computed['fixed']:,.2f}\n"
            f"  Additional Chg   : Rs. {computed['additional']:,.2f}\n"
            f"  Electricity Duty : Rs. {computed['duty']:,.2f}\n"
            f"  Sales Tax        : Rs. {computed['salestax']:,.2f}\n"
            f"  KMC Charges      : Rs. {computed['kmc']:,.2f}\n"
            f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            f"💵 *TOTAL BILL : Rs. {computed['total']:,.2f}*\n"
            f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            f"⚠️ Please pay by the due date to avoid disconnection.\n"
            f"📞 For queries contact: *{COMPANY_NAME}*"
        )
        return plain, wa

    def save_and_print():
        try:
            _do_save_and_print()
        except Exception as _e:
            messagebox.showerror("Unexpected Error",
                f"An error occurred:\n\n{str(_e)}\n\nPlease check all fields and try again.")

    def _do_save_and_print():
        if not computed:
            messagebox.showwarning("Warning", "Please calculate the bill first.")
            return

        cid   = cust_id_ent.get().strip()
        month = month_ent.get().strip()
        acc   = acc_no_ent.get().strip()
        meter = meter_no_ent.get().strip()

        if not cid or not month:
            messagebox.showwarning("Warning", "Please fill Customer ID and Month/Year.")
            return

        # Fetch customer name & phone from DB
        cust_name  = "Unknown"
        prop_title = "Unknown"
        db_phone   = ""
        try:
            wb = load_workbook(DB_FILE)
            ws_cust = wb["Customers"]
            for row in ws_cust.iter_rows(min_row=2, values_only=True):
                if str(row[0]) == cid:
                    cust_name  = row[1]
                    db_phone   = str(row[2])
                    prop_title = row[5]
                    break
        except Exception:
            pass

        # Use manually entered phone if provided, else DB phone
        phone_override = cust_phone_ent.get().strip()
        send_phone = phone_override if phone_override else db_phone

        bill_txt, wa_msg = _get_bill_texts(cid, month, acc, meter, cust_name, prop_title)

        # Save to Excel
        try:
            wb = load_workbook(DB_FILE)
            if "Electricity" not in wb.sheetnames:
                ws_el = wb.create_sheet("Electricity")
                ws_el.append(["Rec_ID", "Cust_ID", "Customer", "Property", "Month",
                              "Prev_Read", "Curr_Read", "Units", "Energy_Charge",
                              "Fixed", "Additional", "Duty", "SalesTax", "KMC",
                              "Total_Bill", "KE_Acc", "Meter_No"])
            else:
                ws_el = wb["Electricity"]
            ws_el.append([
                ws_el.max_row, cid, cust_name, prop_title, month,
                computed["prev"], computed["curr"], computed["units"],
                computed["energy"], computed["fixed"], computed["additional"],
                computed["duty"], computed["salestax"], computed["kmc"],
                computed["total"], acc, meter
            ])
            wb.save(DB_FILE)
        except Exception as e:
            messagebox.showerror("DB Error", f"Could not save to database: {e}")
            return

        # Save Word document
        try:
            doc = Document()
            doc.add_heading(f"{COMPANY_NAME} — Electricity Bill", 0)
            doc.add_paragraph(bill_txt)
            safe_name  = sanitize_filename(cust_name)
            safe_month = sanitize_filename(month)
            fname = os.path.join(_BASE_DIR, f"ElecBill_{safe_name}_{safe_month}.docx")
            doc.save(fname)
        except Exception as doc_err:
            messagebox.showerror("Document Error", f"Could not save Word file:\n{doc_err}")
            return

        # ── Preview Window ──────────────────────────────────────────
        preview_win = tk.Toplevel()
        preview_win.title(f"Bill — {cust_name} — {month}")
        preview_win.geometry("600x560")
        preview_win.configure(bg="white")
        preview_win.resizable(False, False)

        tk.Label(preview_win, text=f"⚡  Electricity Bill — {month}",
                 font=("Segoe UI", 13, "bold"),
                 bg=SIDEBAR_BG, fg=ACCENT).pack(fill="x", ipady=8)

        txt = scrolledtext.ScrolledText(preview_win, font=("Courier New", 10),
                                        width=64, height=20, bd=0, bg="#f9f9f9")
        txt.pack(padx=12, pady=10, fill="both", expand=True)
        txt.insert("1.0", bill_txt)
        txt.config(state="disabled")

        btn_row = tk.Frame(preview_win, bg="white")
        btn_row.pack(pady=10)

        def send_whatsapp_bill():
            """Send professional electricity bill via WhatsApp."""
            num = send_phone.replace(" ", "")
            if not num:
                messagebox.showwarning("No Number",
                    "Enter customer phone in 'Customer WhatsApp No' field or ensure it's in the database.")
                return
            if not num.startswith("+"):
                num = "+92" + num.lstrip("0") if len(num) >= 10 else "+" + num
            try:
                import pywhatkit as kit
                preview_win.withdraw()
                kit.sendwhatmsg_instantly(num, wa_msg, wait_time=15,
                                          tab_close=False, close_time=10)
                time.sleep(12)
                preview_win.deiconify()
                preview_win.lift()
                messagebox.showinfo("Sent ✓", f"Bill sent to {num} via WhatsApp!")
            except Exception as e:
                preview_win.deiconify()
                messagebox.showerror("WhatsApp Error", str(e))

        tk.Button(btn_row, text="📱  Send WhatsApp", command=send_whatsapp_bill,
                  bg=WA_GREEN, fg="white", font=("Segoe UI", 10, "bold"),
                  width=17, height=2, relief="flat", cursor="hand2").pack(side="left", padx=6)
        tk.Button(btn_row, text="📧  Send Email", bg=BTN_BLUE, fg="white",
                  font=("Segoe UI", 10, "bold"), width=14, height=2, relief="flat",
                  command=lambda: Email_Gui(R_frame, bill_txt)).pack(side="left", padx=6)
        tk.Button(btn_row, text="✖  Close", bg="#6c757d", fg="white",
                  font=("Segoe UI", 10), width=10, height=2, relief="flat",
                  command=preview_win.destroy).pack(side="left", padx=6)

        messagebox.showinfo("Saved ✓", f"Bill saved!\nFile: {fname}\nDatabase: Updated ✓")

    def view_electricity_history():
        try:
            wb = load_workbook(DB_FILE)
            if "Electricity" not in wb.sheetnames:
                messagebox.showinfo("Info", "No electricity records found yet.")
                return
            ws_el = wb["Electricity"]
            data = list(ws_el.iter_rows(values_only=True))
        except Exception as e:
            messagebox.showerror("Error", str(e))
            return

        hist_win = tk.Toplevel()
        hist_win.title("Electricity Bill History")
        hist_win.geometry("900x400")
        hist_win.configure(bg="white")

        tk.Label(hist_win, text="Electricity Bill History", font=("Segoe UI", 13, "bold"),
                 bg="#1c2e4a", fg="white").pack(fill="x")

        cols = data[0] if data else []
        tree = ttk.Treeview(hist_win, columns=cols, show="headings", height=16)
        for col in cols:
            tree.heading(col, text=col)
            tree.column(col, width=90, anchor="center")
        for row in data[1:]:
            tree.insert("", "end", values=row)
        scroll = ttk.Scrollbar(hist_win, orient="horizontal", command=tree.xview)
        tree.configure(xscrollcommand=scroll.set)
        tree.pack(fill="both", expand=True, padx=5, pady=5)
        scroll.pack(fill="x")

    # --- Buttons ---
    btn_frame = tk.Frame(outer, bg=CARD_BG)
    btn_frame.pack(pady=15)

    tk.Button(btn_frame, text="⚡  Calculate", command=calculate,
              bg=ACCENT, fg="white", font=("Segoe UI", 11, "bold"),
              width=16, height=2, relief="flat", cursor="hand2").pack(side="left", padx=10)

    tk.Button(btn_frame, text="💾  Save & Print Bill", command=save_and_print,
              bg=BTN_GREEN, fg="white", font=("Segoe UI", 11, "bold"),
              width=18, height=2, relief="flat", cursor="hand2").pack(side="left", padx=10)

    tk.Button(btn_frame, text="📋  View History", command=view_electricity_history,
              bg=BTN_BLUE, fg="white", font=("Segoe UI", 11, "bold"),
              width=16, height=2, relief="flat", cursor="hand2").pack(side="left", padx=10)

    tk.Button(btn_frame, text="← Back", command=lambda: img(R_frame),
              bg="#6c757d", fg="white", font=("Segoe UI", 11),
              width=10, height=2, relief="flat", cursor="hand2").pack(side="left", padx=10)


# --- WHATSAPP / EMAIL / RECEIPT (unchanged) ---

def whatsapp_notice(num, R_frame, name, prop_title, window):
    for widget in R_frame.winfo_children():
        widget.destroy()

    whatsapp_frame = tk.Frame(R_frame, bg="white")
    whatsapp_frame.pack(expand=True, fill="both")

    msg_body = (
        f"        {COMPANY_NAME} - BOOKING CONFIRMATION\n"
        f"To: {name}\n"
        f"Subject: Welcome to Your New Home - {prop_title}\n\n"
        f"Dear {name},\n\n"
        f"Congratulations! Your booking for {prop_title} has been confirmed.\n"
        f"We are excited to have you as our customer.\n\n"
        f"Please visit our office tomorrow at 10:00 AM for the keys and\n"
        f"lease agreement finalization.\n\n"
        f"Best Regards,\n"
        f"Management Team\n"
        f"{COMPANY_NAME}\n"
    )

    body_text = tk.Text(whatsapp_frame, height=18, width=65, font=("Courier", 10), bd=2, relief="groove")
    body_text.pack(expand=True, fill="both", padx=10, pady=10)
    body_text.insert("1.0", msg_body)

    def send_process():
        doc = Document()
        doc.add_heading(f"{COMPANY_NAME} - Booking Confirmation", 0)
        doc.add_paragraph(body_text.get("1.0", tk.END))
        file_path = f"{name}_Booking_Confirmation.docx"
        doc.save(file_path)

        cell_num = str(num).replace(" ", "")
        if not cell_num.startswith("+"):
            cell_num = "+92" + cell_num.lstrip("0") if len(cell_num) >= 10 else "+" + cell_num

        wa_msg = (
            f"* {COMPANY_NAME} Notice *\n\n"
            f"Assalam-o-Alaikum *{name}*,\n"
            f"Congratulations!\n"
            f"Your booking for *{prop_title}* is confirmed.\n"
            f"Please check the attached confirmation letter.\n\n"
            f"Thank You!\n"
            f"*Management - {COMPANY_NAME}*"
        )

        try:
            import pywhatkit as kit
            kit.sendwhatmsg_instantly(cell_num, wa_msg, wait_time=15, tab_close=False, close_time=10)
            time.sleep(12)
            abs_path = os.path.abspath(file_path)
            cmd = f'Powershell -Command "Set-Clipboard -Path \'{abs_path}\'"'
            subprocess.run(cmd, shell=True, check=True)
            time.sleep(2)
            pyautogui.hotkey("ctrl", "v")
            time.sleep(5)
            pyautogui.press("enter")
            window.lift()
            window.focus_force()
            messagebox.showinfo("Success", "WhatsApp message and file sent!")
            img(R_frame)
        except Exception as e:
            messagebox.showerror("Error", f"WhatsApp failed: {e}")

    tk.Button(whatsapp_frame, text=" Send via WhatsApp", command=send_process,
              bg="#25D366", fg="white", height=2).pack(pady=10)


def Email_Gui(R_frame, content):
    for widget in R_frame.winfo_children():
        widget.destroy()

    email_frame = tk.Frame(R_frame, bg="white")
    email_frame.pack(expand=True, fill="both")

    tk.Label(email_frame, text="Customer Email Address:", bg="white").pack(pady=(20, 5))
    rec_entry = tk.Entry(email_frame, width=45)
    rec_entry.pack(pady=5)
    rec_entry.focus_set()

    tk.Label(email_frame, text="Message Body:", bg="white").pack(pady=5)
    body_text = tk.Text(email_frame, height=15, width=60)
    body_text.pack(pady=5)
    body_text.insert("1.0", content)

    def send_email():
        SENDER_EMAIL = "owaisenterprise9105@gmail.com"
        APP_PASSWORD  = "chjk repr qafp ckuq"
        recipient = rec_entry.get()
        subject   = f"Monthly Statement - {COMPANY_NAME}"
        msg = EmailMessage()
        msg.set_content(body_text.get("1.0", tk.END))
        msg['Subject'] = subject
        msg['From']    = SENDER_EMAIL
        msg['To']      = recipient
        try:
            with smtplib.SMTP_SSL("smtp.gmail.com", 465) as server:
                server.login(SENDER_EMAIL, APP_PASSWORD)
                server.send_message(msg)
                messagebox.showinfo("Success", "Email sent successfully!")
                img(R_frame)
        except Exception as e:
            messagebox.showerror("Error", f"Email failed: {e}")

    tk.Button(email_frame, text=" Send Email", command=send_email,
              bg="#007bff", fg="white", height=2).pack(pady=10)


def generate_receipt_gui(R_frame):
    for widget in R_frame.winfo_children():
        widget.destroy()

    rpt_frame = tk.Frame(R_frame, bg="white")
    rpt_frame.pack(expand=True)

    tk.Label(rpt_frame, text=COMPANY_NAME, font=("Arial", 18, "bold"), bg="white").pack(pady=10)
    tk.Label(rpt_frame, text="Enter Customer ID", bg="white").pack(pady=5)
    srch_id = tk.Entry(rpt_frame)
    srch_id.pack(pady=5)

    def process_gen():
        cust_id_raw = srch_id.get()
        if not cust_id_raw.isdigit():
            messagebox.showerror("Error", "Please enter a numeric ID.")
            return

        cust_id = int(cust_id_raw)
        wb = load_workbook(DB_FILE)
        ws_cust = wb["Customers"]
        ws_pay  = wb["Payments"]

        found_cust = None
        for row in ws_cust.iter_rows(min_row=2, values_only=True):
            if row[0] == cust_id:
                found_cust = row
                break

        if not found_cust:
            messagebox.showinfo("Not Found", "Customer record not found.")
            return

        latest_pay = "No payment recorded yet."
        for row in reversed(list(ws_pay.iter_rows(min_row=2, values_only=True))):
            if row[1] == cust_id:
                latest_pay = f"Amount: {row[3]} | Month: {row[4]} | Status: {row[5]}"
                break

        # Latest electricity bill
        elec_info = "No electricity bill recorded yet."
        if "Electricity" in wb.sheetnames:
            ws_el = wb["Electricity"]
            for row in reversed(list(ws_el.iter_rows(min_row=2, values_only=True))):
                if str(row[1]) == str(cust_id):
                    try:
                        total_val = float(row[14]) if row[14] is not None else 0.0
                        elec_info = (f"Month: {row[4]} | Units: {row[7]} kWh | "
                                     f"Bill: Rs. {total_val:,.2f}")
                    except (TypeError, ValueError):
                        elec_info = f"Month: {row[4]} | Units: {row[7]} kWh"
                    break

        receipt_text = (
            f"\t\t {COMPANY_NAME} — RENT RECEIPT \n\n"
            f"Customer ID    : {found_cust[0]}\n"
            f"Name           : {found_cust[1]}\n"
            f"Contact        : {found_cust[2]}\n"
            f"Property Type  : {found_cust[5]}\n"
            f"Duration       : {found_cust[4]} Months\n\n"
            f"LATEST PAYMENT STATUS:\n"
            f"{latest_pay}\n\n"
            f"LATEST ELECTRICITY BILL:\n"
            f"{elec_info}\n\n"
            f"Thank you for Choosing {COMPANY_NAME}!\n"
        )

        doc = Document()
        doc.add_heading(f"{COMPANY_NAME} Statement", 0)
        doc.add_paragraph(receipt_text)
        doc.save(f"Receipt_{found_cust[1]}.docx")

        for widget in rpt_frame.winfo_children():
            widget.destroy()
        view = scrolledtext.ScrolledText(rpt_frame, width=70, height=15)
        view.pack(padx=10, pady=10)
        view.insert("1.0", receipt_text)

        tk.Button(rpt_frame, text="Send Email",
                  command=lambda: Email_Gui(R_frame, receipt_text)).pack(side="left", padx=20, pady=10)
        tk.Button(rpt_frame, text="Back",
                  command=lambda: img(R_frame)).pack(side="right", padx=20, pady=10)
        messagebox.showinfo("Success", "Receipt generated and saved!")

    tk.Button(rpt_frame, text="Generate", command=process_gen, bg="green", fg="white").pack(pady=10)


def payment_entry_gui(R_frame):
    for widget in R_frame.winfo_children():
        widget.destroy()

    pay_frame = tk.Frame(R_frame, bg="white")
    pay_frame.pack(pady=20)

    tk.Label(pay_frame, text="Payment Ledger", font=("Arial", 16, "bold"), bg="white").grid(row=0, columnspan=2, pady=10)

    fields = ["Customer ID", "Customer Name", "Rent Amount", "Utility Charges", "Month/Year"]
    entries = {}
    for i, field in enumerate(fields):
        tk.Label(pay_frame, text=field+":", bg="white").grid(row=i+1, column=0, sticky="w", padx=10, pady=5)
        ent = tk.Entry(pay_frame, width=30)
        ent.grid(row=i+1, column=1, padx=10, pady=5)
        entries[field] = ent

    def save_payment():
        try:
            cid   = entries["Customer ID"].get()
            name  = entries["Customer Name"].get()
            rent  = float(entries["Rent Amount"].get())
            util  = float(entries["Utility Charges"].get())
            month = entries["Month/Year"].get()
            total = rent + util
            wb = load_workbook(DB_FILE)
            ws = wb["Payments"]
            ws.append([ws.max_row, cid, name, total, month, "Paid"])
            wb.save(DB_FILE)
            messagebox.showinfo("Success", f"Payment of {total} recorded for {name}")
            img(R_frame)
        except Exception:
            messagebox.showerror("Error", "Invalid inputs. Ensure numbers are correct.")

    tk.Button(pay_frame, text="Record Payment", command=save_payment, bg="green", fg="white").grid(row=6, column=1, pady=20)
    tk.Button(pay_frame, text="Back", command=lambda: img(R_frame)).grid(row=6, column=0, pady=20)


def view_all_gui(R_frame):
    for widget in R_frame.winfo_children():
        widget.destroy()
    v_frame = tk.Frame(R_frame, bg="white")
    v_frame.pack(fill="both", expand=True)

    wb   = load_workbook(DB_FILE)
    ws   = wb["Customers"]
    data = list(ws.iter_rows(values_only=True))

    if len(data) < 2:
        tk.Label(v_frame, text="No Customers registered yet.", bg="white").pack(pady=20)
        return

    tree = ttk.Treeview(v_frame, columns=data[0], show="headings")
    for col in data[0]:
        tree.heading(col, text=col)
        tree.column(col, width=100)
    for row in data[1:]:
        tree.insert("", "end", values=row)
    tree.pack(fill="both", expand=True)


def property_creation_gui(R_frame):
    for widget in R_frame.winfo_children():
        widget.destroy()
    p_frame = tk.Frame(R_frame, bg="white")
    p_frame.pack(pady=10)

    tk.Label(p_frame, text="Property Listing", font=("Arial", 16, "bold"), bg="white").grid(row=0, columnspan=2, pady=10)

    lbls = ["Property Title", "Address", "Rent Price", "Initial Vacancies"]
    ents = {}
    for i, l in enumerate(lbls):
        tk.Label(p_frame, text=l+":", bg="white").grid(row=i+1, column=0, sticky="w", padx=10, pady=5)
        e = tk.Entry(p_frame, width=30)
        e.grid(row=i+1, column=1, padx=10, pady=5)
        ents[l] = e

    amenities = ["WiFi", "Security", "Parking", "Furnished", "Kitchen", "AC", "Electricity", "Gas"]
    vars = {}
    row_idx = 5
    for i, am in enumerate(amenities):
        var = tk.IntVar()
        tk.Checkbutton(p_frame, text=am, variable=var, bg="white").grid(
            row=row_idx + (i//2), column=i%2, sticky="w", padx=10)
        vars[am] = var

    def save_prop():
        try:
            title = ents["Property Title"].get()
            addr  = ents["Address"].get()
            rent  = ents["Rent Price"].get()
            vac   = int(ents["Initial Vacancies"].get())
            am_vals = [vars[a].get() for a in amenities]
            wb = load_workbook(DB_FILE)
            ws = wb["Properties"]
            ws.append([ws.max_row, title, addr, rent, vac] + am_vals)
            wb.save(DB_FILE)
            messagebox.showinfo("Success", "Property Added Successfully!")
            img(R_frame)
        except Exception as e:
            messagebox.showerror("Error", f"Failed: {e}")

    tk.Button(p_frame, text="Submit", command=save_prop, bg="blue", fg="white").grid(row=row_idx+5, column=1, pady=10)
    tk.Button(p_frame, text="Back",   command=lambda: img(R_frame)).grid(row=row_idx+5, column=0, pady=10)


def customer_app_gui(R_frame, window):
    for widget in R_frame.winfo_children():
        widget.destroy()
    c_frame = tk.Frame(R_frame, bg="white")
    c_frame.pack(pady=10)

    tk.Label(c_frame, text="Customer Registration", font=("Arial", 16, "bold"), bg="white").grid(row=0, columnspan=2, pady=10)

    lbls = ["Full Name", "Contact Number", "Stay Duration (Months)"]
    ents = {}
    for i, l in enumerate(lbls):
        tk.Label(c_frame, text=l+":", bg="white").grid(row=i+1, column=0, sticky="w", padx=10, pady=5)
        e = tk.Entry(c_frame, width=30)
        e.grid(row=i+1, column=1, padx=10, pady=5)
        ents[l] = e

    tk.Label(c_frame, text="Select Amenities Required:", bg="white", font=("Arial", 10, "bold")).grid(row=4, columnspan=2, pady=5)
    amenities = ["WiFi", "Security", "Parking", "Furnished", "Kitchen", "AC", "Electricity", "Gas"]
    vars = {}
    for i, am in enumerate(amenities):
        var = tk.IntVar()
        tk.Checkbutton(c_frame, text=am, variable=var, bg="white").grid(
            row=5 + (i//2), column=i%2, sticky="w", padx=10)
        vars[am] = var

    def register():
        name = ents["Full Name"].get()
        num  = ents["Contact Number"].get()
        dur  = ents["Stay Duration (Months)"].get()
        user_ams = [vars[a].get() for a in amenities]

        wb = load_workbook(DB_FILE)
        ws_prop = wb["Properties"]

        selected_prop = None
        for row_idx, row in enumerate(ws_prop.iter_rows(min_row=2, values_only=True), 2):
            vac = int(row[4] or 0)
            if vac > 0:
                prop_ams = row[5:]
                match = True
                for i in range(len(user_ams)):
                    if user_ams[i] == 1 and prop_ams[i] == 0:
                        match = False; break
                if match:
                    ws_prop.cell(row=row_idx, column=5).value = vac - 1
                    selected_prop = row
                    break

        if selected_prop:
            ws_cust = wb["Customers"]
            cid = ws_cust.max_row
            ws_cust.append([cid, name, num, selected_prop[0], dur, selected_prop[1]])
            wb.save(DB_FILE)
            messagebox.showinfo("Success", f"Booked Successfully! Your property: {selected_prop[1]}")
            whatsapp_notice(num, R_frame, name, selected_prop[1], window)
        else:
            messagebox.showinfo("Sorry", "No matching property with vacancy available.")
            img(R_frame)

    tk.Button(c_frame, text="Register", command=register, bg="green", fg="white").grid(row=10, column=1, pady=10)
    tk.Button(c_frame, text="Back",     command=lambda: img(R_frame)).grid(row=10, column=0, pady=10)


# --- NAVIGATION & APP WRAPPER ---

def admin_menu(menu_frame, main_frame, window):
    for widget in menu_frame.winfo_children():
        widget.destroy()
    menu_frame.config(bg="#1c2e4a")

    tk.Label(menu_frame, text="ADMIN PANEL", font=("Arial", 14, "bold"),
             bg="#1c2e4a", fg="white").pack(pady=20)

    btns = [
        ("🏠  Add Property",      lambda: property_creation_gui(main_frame)),
        ("👥  View Customers",    lambda: view_all_gui(main_frame)),
        ("💰  Record Payment",    lambda: payment_entry_gui(main_frame)),
        ("📄  Gen Receipt",       lambda: generate_receipt_gui(main_frame)),
        ("⚡  Electricity Bill",  lambda: electricity_bill_gui(main_frame)),
        ("🏠  Main Menu",         lambda: main_menu(menu_frame, main_frame, window)),
    ]

    for txt, cmd in btns:
        tk.Button(menu_frame, text=txt, font=("Arial", 11), width=18, command=cmd,
                  bg="#2a3f5a", fg="white", activebackground="#3d5a80",
                  anchor="w", padx=10).pack(pady=5)


def admin_login_gui(menu_frame, main_frame, window):
    for widget in main_frame.winfo_children():
        widget.destroy()
    l_frame = tk.Frame(main_frame, bg="white")
    l_frame.pack(expand=True)

    tk.Label(l_frame, text="Security Access", font=("Arial", 16, "bold"), bg="white").pack(pady=10)
    tk.Label(l_frame, text="Enter Password:", bg="white").pack()
    pw_ent = tk.Entry(l_frame, show="*")
    pw_ent.pack(pady=10)

    def check():
        if pw_ent.get() == "1234":
            admin_menu(menu_frame, main_frame, window)
            img(main_frame)
        else:
            messagebox.showerror("Error", "Invalid Password")

    tk.Button(l_frame, text="Login", command=check, bg="#007bff", fg="white", width=10).pack()


def main_menu(menu_frame, main_frame, window):
    for widget in menu_frame.winfo_children():
        widget.destroy()
    menu_frame.config(bg="#f8f9fa")

    try:
        logo_img = Image.open(LOGO_IMAGE_PATH).resize((150, 150), Image.Resampling.LANCZOS)
        logo_tk  = ImageTk.PhotoImage(logo_img)
        logo_lbl = tk.Label(menu_frame, image=logo_tk, bg="#f8f9fa")
        logo_lbl.image = logo_tk
        logo_lbl.pack(pady=20)
    except:
        pass

    tk.Label(menu_frame, text="PORTAL", font=("Arial", 16, "bold"), bg="#f8f9fa").pack(pady=10)
    tk.Button(menu_frame, text="Admin Login",   font=("Arial", 12), width=15,
              command=lambda: admin_login_gui(menu_frame, main_frame, window)).pack(pady=10)
    tk.Button(menu_frame, text="Customer App",  font=("Arial", 12), width=15,
              command=lambda: customer_app_gui(main_frame, window)).pack(pady=10)
    tk.Button(menu_frame, text="Exit",          font=("Arial", 12), width=15,
              command=exit).pack(pady=10)

    img(main_frame)


# --- START APP ---
def start_app():
    root = tk.Tk()
    root.title(COMPANY_NAME)
    root.geometry("1100x650")

    sidebar = tk.Frame(root, width=250, bg="#f8f9fa", borderwidth=1, relief="flat")
    sidebar.pack(side="left", fill="y")

    main_frame = tk.Frame(root, bg="white")
    main_frame.pack(side="right", fill="both", expand=True)

    main_menu(sidebar, main_frame, root)
    root.mainloop()


if __name__ == "__main__":
    start_app()
