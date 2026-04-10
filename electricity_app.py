import os
import re
import time
import subprocess
import smtplib
from email.message import EmailMessage
from datetime import datetime
import tkinter as tk
from tkinter import ttk, messagebox, filedialog, scrolledtext, simpledialog
from PIL import Image, ImageTk
from openpyxl import load_workbook, Workbook
from docx import Document

try:
    import pytesseract
    _tess_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    if os.path.exists(_tess_path):
        pytesseract.pytesseract.tesseract_cmd = _tess_path
        OCR_AVAILABLE = True
    else:
        OCR_AVAILABLE = False
except ImportError:
    OCR_AVAILABLE = False
except Exception:
    OCR_AVAILABLE = False

# --- CONFIGURATION ---
_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_FILE = os.path.join(_BASE_DIR, "Rent_Management.xlsx")
COMPANY_NAME = "Owais Rent Management"

KE_SLABS = [
    (100, 10.54),
    (200, 13.01),
    (float('inf'), 13.01),
]
FIXED_CHARGE     = 600.00
ADDITIONAL_CHG   = 169.00
ELEC_DUTY        = 19.91
SALES_TAX        = 350.53
KMC_CHG          = 20.00

def sanitize_filename(name: str) -> str:
    return re.sub(r'[\\/:*?"<>|]', '_', name).strip() or "Unknown"

def calc_ke_bill(units, additional=ADDITIONAL_CHG):
    energy = 0.0
    remaining = units
    prev_limit = 0
    for limit, rate in KE_SLABS:
        if remaining <= 0:
            break
        slab_units = min(remaining, limit - prev_limit)
        energy += slab_units * rate
        remaining -= slab_units
        prev_limit = limit
    total = energy + FIXED_CHARGE + additional + ELEC_DUTY + SALES_TAX + KMC_CHG
    return (
        round(energy, 2), round(FIXED_CHARGE, 2), round(additional, 2),
        round(ELEC_DUTY, 2), round(SALES_TAX, 2), round(KMC_CHG, 2), round(total, 2)
    )

class ElectricityStandaloneApp:
    def __init__(self, root):
        self.root = root
        self.root.title("⚡ Electricity Bill Calculator - Standalone")
        self.root.geometry("1050x700")
        self.root.configure(bg="#f4f6f9")
        self.image_path = None
        self.computed = {}
        # Fixed customer defaults
        self.customer_data = {
            'name':  'MEER MUZAFFARALI',
            'phone': '',
            'prop':  'Muzaffer'
        }
        
        self.build_ui()
        self.update_clock()

    def update_clock(self):
        now = datetime.now().strftime("%I:%M:%S %p  |  %d-%b-%Y")
        self.time_lbl.config(text=now)
        self.root.after(1000, self.update_clock)

    def build_ui(self):
        # Header
        header = tk.Frame(self.root, bg="#1c2e4a", pady=15)
        header.pack(fill="x")
        
        tk.Label(header, text="⚡ Smart Electricity Calculator", font=("Segoe UI", 18, "bold"), bg="#1c2e4a", fg="#f0a500").pack()
        self.time_lbl = tk.Label(header, text="", font=("Segoe UI", 11), bg="#1c2e4a", fg="white")
        self.time_lbl.pack()

        # Sync entry fields → customer_data on fetch
        # Main Body Frame
        content_frame = tk.Frame(self.root, bg="#f4f6f9")
        content_frame.pack(fill="both", expand=True, padx=20, pady=10)

        # Left Column - Inputs
        left_frame = tk.Frame(content_frame, bg="white", relief="groove", bd=1, padx=20, pady=15)
        left_frame.pack(side="left", fill="both", expand=True, padx=(0, 10))

        # Right Column - Image & Results
        right_frame = tk.Frame(content_frame, bg="white", relief="groove", bd=1, padx=20, pady=15)
        right_frame.pack(side="right", fill="both", expand=True)

        # --- LEFT FRAME (Inputs) ---
        tk.Label(left_frame, text="1. Customer Details", font=("Segoe UI", 12, "bold"), bg="white", fg="#1a1a2e").grid(row=0, column=0, columnspan=3, sticky="w", pady=(0, 10))

        tk.Label(left_frame, text="Cust ID:", bg="white", font=("Segoe UI", 10)).grid(row=1, column=0, sticky="w", pady=5)
        self.ent_cid = tk.Entry(left_frame, font=("Segoe UI", 11), width=14)
        self.ent_cid.insert(0, "LA104056")
        self.ent_cid.grid(row=1, column=1, sticky="w", pady=5)
        tk.Button(left_frame, text="🔍 Fetch", bg="#007bff", fg="white", cursor="hand2", command=self.fetch_customer_record).grid(row=1, column=2, sticky="w", padx=10)

        tk.Label(left_frame, text="Name:", bg="white", font=("Segoe UI", 10)).grid(row=2, column=0, sticky="w", pady=5)
        self.ent_name = tk.Entry(left_frame, font=("Segoe UI", 11), width=22)
        self.ent_name.insert(0, "MEER MUZAFFARALI")
        self.ent_name.grid(row=2, column=1, columnspan=2, sticky="w", pady=5)

        tk.Label(left_frame, text="Property:", bg="white", font=("Segoe UI", 10)).grid(row=3, column=0, sticky="w", pady=5)
        self.ent_prop = tk.Entry(left_frame, font=("Segoe UI", 11), width=22)
        self.ent_prop.insert(0, "Muzaffer")
        self.ent_prop.grid(row=3, column=1, columnspan=2, sticky="w", pady=5)

        self.lbl_cust_info = tk.Label(left_frame, text="", bg="white", font=("Segoe UI", 9, "italic"), fg="#28a745")
        self.lbl_cust_info.grid(row=4, column=0, columnspan=3, sticky="w", pady=(0, 8))

        tk.Label(left_frame, text="2. Meter Readings", font=("Segoe UI", 12, "bold"), bg="white", fg="#1a1a2e").grid(row=5, column=0, columnspan=3, sticky="w", pady=(10, 10))

        tk.Label(left_frame, text="Prev Reading (Auto):", bg="white", font=("Segoe UI", 10)).grid(row=6, column=0, sticky="w", pady=5)
        self.ent_prev = tk.Entry(left_frame, font=("Segoe UI", 11, "bold"), fg="#dc3545", width=12)
        self.ent_prev.grid(row=6, column=1, sticky="w", pady=5)
        self.lbl_prev_date = tk.Label(left_frame, text="(Date: --)", bg="white", font=("Segoe UI", 9), fg="#888")
        self.lbl_prev_date.grid(row=6, column=2, sticky="w", padx=5)

        tk.Label(left_frame, text="Curr Reading (New):", bg="white", font=("Segoe UI", 10)).grid(row=7, column=0, sticky="w", pady=5)
        self.ent_curr = tk.Entry(left_frame, font=("Segoe UI", 11, "bold"), fg="#28a745", width=12)
        self.ent_curr.grid(row=7, column=1, sticky="w", pady=5)

        tk.Label(left_frame, text="Month / Year:", bg="white", font=("Segoe UI", 10)).grid(row=8, column=0, sticky="w", pady=5)
        self.ent_month = tk.Entry(left_frame, font=("Segoe UI", 11), width=12)
        self.ent_month.grid(row=8, column=1, sticky="w", pady=5)
        self.ent_month.insert(0, datetime.now().strftime("%b-%Y"))

        tk.Label(left_frame, text="Additional Chg (Rs):", bg="white", font=("Segoe UI", 10)).grid(row=9, column=0, sticky="w", pady=5)
        self.ent_add_chg = tk.Entry(left_frame, font=("Segoe UI", 11), width=12)
        self.ent_add_chg.insert(0, str(ADDITIONAL_CHG))
        self.ent_add_chg.grid(row=9, column=1, sticky="w", pady=5)

        # Buttons
        tk.Button(left_frame, text="📷 Upload Pic & Auto-Read", bg="#17a2b8", fg="white", font=("Segoe UI", 10, "bold"), width=30, command=self.process_image, pady=5, cursor="hand2").grid(row=10, column=0, columnspan=3, pady=(20, 8))
        tk.Button(left_frame, text="⚡ Calculate Bill", bg="#f0a500", fg="white", font=("Segoe UI", 11, "bold"), width=28, command=self.calculate_bill, pady=5, cursor="hand2").grid(row=11, column=0, columnspan=3, pady=5)
        tk.Button(left_frame, text="💾 Save, Print & Share", bg="#28a745", fg="white", font=("Segoe UI", 11, "bold"), width=28, command=self.save_and_preview, pady=5, cursor="hand2").grid(row=12, column=0, columnspan=3, pady=5)


        # --- RIGHT FRAME (Results & Image) ---
        tk.Label(right_frame, text="Meter Image Preview", font=("Segoe UI", 12, "bold"), bg="white", fg="#1a1a2e").pack(anchor="w")
        self.lbl_img_preview = tk.Label(right_frame, text="[No Image Selected]", bg="#e9ecef", width=45, height=10)
        self.lbl_img_preview.pack(pady=10)

        tk.Label(right_frame, text="Bill Summary Breakdown", font=("Segoe UI", 12, "bold"), bg="white", fg="#1a1a2e").pack(anchor="w", pady=(15, 5))
        
        self.res_vars = {}
        fields = [("Units Consumed:", "units"), ("Energy Charges:", "energy"), ("Fixed Charge:", "fixed"), 
                  ("Additional:", "additional"), ("Sales Tax & Duty:", "tax"), ("KMC Charges:", "kmc"), ("TOTAL BILL:", "total")]
        
        res_grid = tk.Frame(right_frame, bg="white")
        res_grid.pack(fill="x")
        
        for i, (lbl, key) in enumerate(fields):
            tk.Label(res_grid, text=lbl, bg="white", font=("Segoe UI", 10)).grid(row=i, column=0, sticky="w", pady=4)
            var = tk.StringVar(value="---")
            self.res_vars[key] = var
            bold = "bold" if key in ["units", "total"] else "normal"
            color = "#dc3545" if key == "total" else "#28a745" if key == "units" else "black"
            tk.Label(res_grid, textvariable=var, bg="white", font=("Segoe UI", 11, bold), fg=color).grid(row=i, column=1, sticky="e", pady=4, padx=20)

    def fetch_customer_record(self):
        cid = self.ent_cid.get().strip()
        if not cid:
            messagebox.showwarning("Missing ID", "Please enter a Customer ID to search.")
            return

        try:
            wb = load_workbook(DB_FILE, data_only=True)
            if "Customers" not in wb.sheetnames:
                messagebox.showerror("No Data", "Customers sheet not found in Database.")
                return
                
            ws_cust = wb["Customers"]
            found = False
            for row in ws_cust.iter_rows(min_row=2, values_only=True):
                if str(row[0]) == cid:
                    self.customer_data['name'] = str(row[1]) if row[1] else "Unknown"
                    self.customer_data['phone'] = str(row[2]) if row[2] else ""
                    self.customer_data['prop'] = str(row[5]) if len(row) > 5 and row[5] else "Unknown"
                    # Update entry fields
                    self.ent_name.delete(0, tk.END)
                    self.ent_name.insert(0, self.customer_data['name'])
                    self.ent_prop.delete(0, tk.END)
                    self.ent_prop.insert(0, self.customer_data['prop'])
                    self.lbl_cust_info.config(text=f"✅ Record found: {self.customer_data['name']} | {self.customer_data['prop']}")
                    found = True
                    break
            
            if not found:
                messagebox.showerror("Not Found", "Customer ID not found in database.")
                return

            # Now fetch last reading
            self.ent_prev.delete(0, tk.END)
            self.lbl_prev_date.config(text="(Date: --)")
            
            if "Electricity" in wb.sheetnames:
                ws_el = wb["Electricity"]
                last_reading = None
                last_month = None
                
                # Reverse iteration to find latest
                for row in reversed(list(ws_el.iter_rows(min_row=2, values_only=True))):
                    if str(row[1]) == cid:
                        last_reading = row[6]   # Curr_Read
                        last_month   = row[4]   # Month
                        # Date_Issued is at index 17 (if saved by new version)
                        last_date    = row[17] if len(row) > 17 and row[17] else last_month
                        break

                if last_reading is not None:
                    self.ent_prev.insert(0, str(last_reading))
                    self.lbl_prev_date.config(text=f"📅 {last_date}")
                    messagebox.showinfo(
                        "Auto-Fetched ✓",
                        f"Previous Reading: {last_reading} kWh\nRecorded on: {last_date}"
                    )
                else:
                    messagebox.showinfo("Info", "No previous electricity records found for this customer. Enter manually.")

        except Exception as e:
            messagebox.showerror("Database Error", f"Error accessing database:\n{e}")

    def process_image(self):
        file_path = filedialog.askopenfilename(filetypes=[("Image Files", "*.jpg;*.jpeg;*.png")])
        if not file_path:
            return

        self.image_path = file_path
        
        # Display Preview
        try:
            img = Image.open(file_path)
            img.thumbnail((300, 200))
            img_tk = ImageTk.PhotoImage(img)
            self.lbl_img_preview.config(image=img_tk, text="")
            self.lbl_img_preview.image = img_tk
        except Exception:
            self.lbl_img_preview.config(text="[Error Loading Image]")

        # Perform OCR — silently skip if tesseract not installed
        if not OCR_AVAILABLE:
            self.lbl_cust_info.config(
                text="📷 Image loaded — Please read meter and type in 'Curr Reading' above",
                fg="#f0a500"
            )
            return

        try:
            custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789'
            text = pytesseract.image_to_string(Image.open(file_path), config=custom_config)
            numbers = re.findall(r'\d+', text)
            if numbers:
                largest_num = max(numbers, key=len)
                self.ent_curr.delete(0, tk.END)
                self.ent_curr.insert(0, largest_num)
                self.lbl_cust_info.config(
                    text=f"✅ OCR Reading: {largest_num} — Please verify from image above!",
                    fg="#28a745"
                )
            else:
                self.lbl_cust_info.config(
                    text="⚠ OCR: No number detected — type 'Curr Reading' manually",
                    fg="#dc3545"
                )
        except Exception:
            self.lbl_cust_info.config(
                text="📷 Image loaded — Please type 'Curr Reading' manually",
                fg="#f0a500"
            )

    def calculate_bill(self):
        try:
            prev = int(self.ent_prev.get())
            curr = int(self.ent_curr.get())
            add_chg = float(self.ent_add_chg.get())
        except ValueError:
            messagebox.showerror("Input Error", "Please ensure previous reading, current reading, and additional charges are valid numbers.")
            return

        if curr < prev:
            messagebox.showerror("Math Error", "Current reading cannot be less than previous reading.")
            return

        units = curr - prev
        energy, fixed, additional, duty, salestax, kmc, total = calc_ke_bill(units, add_chg)

        self.res_vars["units"].set(f"{units} kWh")
        self.res_vars["energy"].set(f"Rs. {energy:,.2f}")
        self.res_vars["fixed"].set(f"Rs. {fixed:,.2f}")
        self.res_vars["additional"].set(f"Rs. {additional:,.2f}")
        self.res_vars["tax"].set(f"Rs. {(duty + salestax):,.2f}")
        self.res_vars["kmc"].set(f"Rs. {kmc:,.2f}")
        self.res_vars["total"].set(f"Rs. {total:,.2f}")

        self.computed = {
            "prev": prev, "curr": curr, "units": units,
            "energy": energy, "fixed": fixed, "additional": additional,
            "duty": duty, "salestax": salestax, "kmc": kmc, "total": total
        }

    def save_and_preview(self):
        if not self.computed:
            messagebox.showwarning("Wait", "Calculate the bill first!")
            return
            
        cid = self.ent_cid.get().strip()
        month = self.ent_month.get().strip()
        
        if not cid or not month:
            messagebox.showwarning("Wait", "Customer ID and Month cannot be empty.")
            return
            
        # Always read from entry fields (user may have edited them)
        name  = self.ent_name.get().strip() or self.customer_data.get('name', 'Unknown')
        prop  = self.ent_prop.get().strip() or self.customer_data.get('prop', 'Unknown')
        phone = self.customer_data.get('phone', '')

        # Save to DB
        try:
            wb = load_workbook(DB_FILE)
            if "Electricity" not in wb.sheetnames:
                ws_el = wb.create_sheet("Electricity")
                ws_el.append(["Rec_ID", "Cust_ID", "Customer", "Property", "Month",
                              "Prev_Read", "Curr_Read", "Units", "Energy_Charge",
                              "Fixed", "Additional", "Duty", "SalesTax", "KMC",
                              "Total_Bill", "KE_Acc", "Meter_No", "Date_Issued"])
            else:
                ws_el = wb["Electricity"]
                
            ws_el.append([
                ws_el.max_row, cid, name, prop, month,
                self.computed["prev"], self.computed["curr"], self.computed["units"],
                self.computed["energy"], self.computed["fixed"], self.computed["additional"],
                self.computed["duty"], self.computed["salestax"], self.computed["kmc"],
                self.computed["total"], "0400013715589", "SEA78555",
                datetime.now().strftime("%d-%b-%Y %I:%M %p")
            ])
            wb.save(DB_FILE)
        except Exception as e:
            messagebox.showerror("DB Error", f"Could not save to Excel:\n{e}")
            return

        now_str = datetime.now().strftime("%d-%b-%Y %I:%M %p")
        
        plain_txt = (
            f"\n{'='*55}\n"
            f"         {COMPANY_NAME}\n"
            f"          ELECTRICITY BILL RECEIPT\n"
            f"{'='*55}\n"
            f"  Generated On: {now_str}\n"
            f"  Month       : {month}\n"
            f"  Customer ID : {cid}\n"
            f"  Name        : {name}\n"
            f"  Property    : {prop}\n"
            f"{'-'*55}\n"
            f"  METER READINGS\n"
            f"  Previous Reading : {self.computed['prev']} kWh\n"
            f"  Current Reading  : {self.computed['curr']} kWh\n"
            f"  Units Consumed   : {self.computed['units']} kWh\n"
            f"{'-'*55}\n"
            f"  CHARGES\n"
            f"  Energy           : Rs. {self.computed['energy']:,.2f}\n"
            f"  Fixed Charge     : Rs. {self.computed['fixed']:,.2f}\n"
            f"  Additional       : Rs. {self.computed['additional']:,.2f}\n"
            f"  Taxes & Duty     : Rs. {(self.computed['duty'] + self.computed['salestax']):,.2f}\n"
            f"  KMC Charges      : Rs. {self.computed['kmc']:,.2f}\n"
            f"{'-'*55}\n"
            f"  TOTAL BILL       : Rs. {self.computed['total']:,.2f}\n"
            f"{'='*55}\n"
        )
        
        try:
            doc = Document()
            doc.add_heading(f"{COMPANY_NAME} — Electricity Bill", 0)
            doc.add_paragraph(plain_txt)
            fname = os.path.join(_BASE_DIR, f"ElecBill_{sanitize_filename(name)}_{sanitize_filename(month)}.docx")
            doc.save(fname)
            messagebox.showinfo("Saved", f"Bill document saved as:\n{fname}")
        except Exception as e:
            messagebox.showerror("File Error", f"Failed to save document:\n{e}")

        self.open_preview(plain_txt, phone, name, month)

    def open_preview(self, bill_txt, phone, name, month):
        win = tk.Toplevel(self.root)
        win.title(f"Bill Preview - {name}")
        win.geometry("600x600")
        
        tk.Label(win, text=f"Bill Preview: {month}", font=("Segoe UI", 12, "bold"), bg="#1c2e4a", fg="white").pack(fill="x", ipady=5)
        
        txt = scrolledtext.ScrolledText(win, font=("Courier New", 10), width=65, height=22)
        txt.pack(padx=10, pady=10, fill="both", expand=True)
        txt.insert("1.0", bill_txt)
        txt.config(state="disabled")
        
        btn_frame = tk.Frame(win)
        btn_frame.pack(pady=10)
        
        wa_msg = (
            f"⚡ *ELECTRICITY BILL — {COMPANY_NAME}*\n"
            f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            f"📅 *Month:* {month}\n"
            f"👤 *Customer:* {name}\n"
            f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            f"  Previous  : {self.computed['prev']} kWh\n"
            f"  Current   : {self.computed['curr']} kWh\n"
            f"  *Units Used : {self.computed['units']} kWh*\n"
            f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            f"💵 *TOTAL BILL : Rs. {self.computed['total']:,.2f}*\n"
            f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
            f"⌚ Time: {datetime.now().strftime('%d-%b-%Y %I:%M %p')}\n"
            f"⚠️ Please pay by the due date."
        )

        def send_wa():
            target_phone = phone.replace(" ", "")
            if not target_phone:
                target_phone = simpledialog.askstring("Phone Number", "Enter WhatsApp Number (e.g., 03001234567):")
                if not target_phone: return
            
            if not target_phone.startswith("+"):
                target_phone = "+92" + target_phone.lstrip("0") if len(target_phone) >= 10 else "+" + target_phone
                
            try:
                import pywhatkit as kit
                win.withdraw()
                messagebox.showinfo("WhatsApp", "Browser will open shortly to send the message. Please do not touch mouse/keyboard during process.")
                kit.sendwhatmsg_instantly(target_phone, wa_msg, wait_time=15, tab_close=False, close_time=10)
                time.sleep(12)
                win.deiconify()
                messagebox.showinfo("Sent", "Message sent via WhatsApp!")
            except Exception as e:
                win.deiconify()
                messagebox.showerror("Error", f"Failed to send WhatsApp:\n{e}")

        tk.Button(btn_frame, text="📱 Send WhatsApp", bg="#25D366", fg="white", font=("Segoe UI", 10, "bold"), width=15, command=send_wa).pack(side="left", padx=5)
        tk.Button(btn_frame, text="✖ Close", bg="#6c757d", fg="white", font=("Segoe UI", 10), width=15, command=win.destroy).pack(side="left", padx=5)


if __name__ == "__main__":
    root = tk.Tk()
    app = ElectricityStandaloneApp(root)
    root.mainloop()
