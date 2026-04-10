import os
import re
import time
import json
import streamlit as st
import smtplib
from email.message import EmailMessage
from datetime import datetime
from PIL import Image
from openpyxl import load_workbook, Workbook
from docx import Document

try:
    import pytesseract
    if os.name == 'nt':  # Windows environment
        _tess_path = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
        if os.path.exists(_tess_path):
            pytesseract.pytesseract.tesseract_cmd = _tess_path
            OCR_AVAILABLE = True
        else:
            OCR_AVAILABLE = False
    else:
        # Linux/Streamlit Cloud environment
        import shutil
        if shutil.which("tesseract"):
            OCR_AVAILABLE = True
        else:
            OCR_AVAILABLE = False
except ImportError:
    OCR_AVAILABLE = False
except Exception:
    OCR_AVAILABLE = False

# --- CONFIGURATION ---
_BASE_DIR = os.path.dirname(os.path.abspath(__file__))
DB_FILE = os.path.join(_BASE_DIR, "Rent_Management.xlsx")
FIXED_READING_FILE = os.path.join(_BASE_DIR, "fixed_reading.json")
COMPANY_NAME = "Owais Rent Management"

KE_SLABS = [
    (100, 10.54),
    (200, 13.01),
    (float('inf'), 13.01),
]
FIXED_CHARGE     = 600.00
ADDITIONAL_CHG   = 169.00
ELEC_DUTY        = 19.91
SALES_TAX        = 350.53
KMC_CHG          = 20.00

def sanitize_filename(name: str) -> str:
    return re.sub(r'[\\/:*?"<>|]', '_', name).strip() or "Unknown"

def calc_ke_bill(units, additional=ADDITIONAL_CHG):
    energy = 0.0
    remaining = units
    prev_limit = 0
    for limit, rate in KE_SLABS:
        if remaining <= 0:
            break
        slab_units = min(remaining, limit - prev_limit)
        energy += slab_units * rate
        remaining -= slab_units
        prev_limit = limit
    total = energy + FIXED_CHARGE + additional + ELEC_DUTY + SALES_TAX + KMC_CHG
    return (
        round(energy, 2), round(FIXED_CHARGE, 2), round(additional, 2),
        round(ELEC_DUTY, 2), round(SALES_TAX, 2), round(KMC_CHG, 2), round(total, 2)
    )

st.set_page_config(page_title="Electricity Calculator", page_icon="⚡", layout="wide")

# CSS Injection for Premium Look
st.markdown("""
<style>
    .main {
        background-color: #f4f6f9;
        font-family: 'Segoe UI', Tahoma, Geneva, Verdana, sans-serif;
    }
    .stButton>button {
        background-color: #007bff;
        color: white;
        border-radius: 8px;
        padding: 0.5rem 1rem;
        font-weight: bold;
        border: none;
        transition: 0.3s;
    }
    .stButton>button:hover {
        background-color: #0056b3;
    }
    .metric-card {
        background-color: white;
        padding: 15px;
        border-radius: 10px;
        box-shadow: 0 4px 6px rgba(0,0,0,0.1);
        margin-bottom: 10px;
        border-left: 5px solid #007bff;
    }
    .metric-title {
        font-size: 14px;
        color: #6c757d;
        margin-bottom: 5px;
        font-weight: bold;
    }
    .metric-value {
        font-size: 24px;
        font-weight: bold;
        color: #1a1a2e;
    }
    .metric-value.total { color: #dc3545; }
    .metric-value.units { color: #28a745; }
</style>
""", unsafe_allow_html=True)

# Load fixed reading from JSON
fixed_val = 0
fixed_d = "--"
if os.path.exists(FIXED_READING_FILE):
    try:
        with open(FIXED_READING_FILE, "r") as f:
            data = json.load(f)
            fixed_val = data.get("reading", 0)
            fixed_d = data.get("date", "--")
    except Exception:
        pass

# Custom session state variables
if 'customer_data' not in st.session_state:
    st.session_state.customer_data = {'name': 'Irshad', 'prop': 'Meer Muzaffar Ali', 'phone': ''}
if 'computed' not in st.session_state:
    st.session_state.computed = None
if 'prev_reading' not in st.session_state:
    st.session_state.prev_reading = fixed_val
if 'prev_date' not in st.session_state:
    st.session_state.prev_date = fixed_d
if 'ocr_reading' not in st.session_state:
    st.session_state.ocr_reading = 0

# Header
st.markdown("<h2 style='text-align: center; color: #1c2e4a; font-weight: bold;'>⚡ Smart Electricity Calculator</h2>", unsafe_allow_html=True)
st.markdown("""<hr style="height:2px;border-width:0;color:gray;background-color:#007bff; margin-top: 0px;">""", unsafe_allow_html=True)

col1, col2 = st.columns([1.2, 1], gap="large")

with col1:
    st.markdown("### 👤 1. Customer Details")
    
    with st.container():
        c1, c2 = st.columns([3, 1])
        with c1:
            cid_input = st.text_input("Customer ID:", value="LA104056")
        with c2:
            st.markdown("<br>", unsafe_allow_html=True)
            fetch_btn = st.button("🔍 Fetch DB", use_container_width=True)
            
        if fetch_btn:
            try:
                wb = load_workbook(DB_FILE, data_only=True)
                if "Customers" in wb.sheetnames:
                    ws_cust = wb["Customers"]
                    found = False
                    for row in ws_cust.iter_rows(min_row=2, values_only=True):
                        if str(row[0]) == cid_input.strip():
                            st.session_state.customer_data['name'] = str(row[1]) if row[1] else "Unknown"
                            st.session_state.customer_data['phone'] = str(row[2]) if row[2] else ""
                            st.session_state.customer_data['prop'] = str(row[5]) if len(row) > 5 and row[5] else "Unknown"
                            found = True
                            st.success(f"✅ Found: {st.session_state.customer_data['name']} ({st.session_state.customer_data['prop']})")
                            break
                    if not found:
                        st.error("Customer ID not found.")
                
                if "Electricity" in wb.sheetnames:
                    ws_el = wb["Electricity"]
                    last_reading = None
                    last_date = None
                    for row in reversed(list(ws_el.iter_rows(min_row=2, values_only=True))):
                        if str(row[1]) == cid_input.strip():
                            last_reading = row[6]
                            last_month = row[4]
                            last_date = row[17] if len(row) > 17 and row[17] else last_month
                            break
                    if last_reading is not None:
                        st.session_state.prev_reading = int(last_reading)
                        st.session_state.prev_date = last_date
                        st.info(f"📅 Auto-Fetched Prev Reading: **{last_reading} kWh** on **{last_date}**")
            except Exception as e:
                st.error(f"DB Error: {e}")

    c_name_col, c_prop_col = st.columns(2)
    with c_name_col:
        cname = st.text_input("Customer Name:", value=st.session_state.customer_data['name'])
    with c_prop_col:
        cprop = st.text_input("Property:", value=st.session_state.customer_data['prop'])
    
    st.markdown("---")
    st.markdown("### 📟 2. Meter Readings")
    pr_col1, pr_col2 = st.columns([2, 1])
    with pr_col1:
        prev_r = st.number_input(f"Previous Reading (Fixed Date: {st.session_state.prev_date}):", value=int(st.session_state.prev_reading))
    with pr_col2:
        st.markdown("<br>", unsafe_allow_html=True)
        if st.button("📌 Fix this Reading", use_container_width=True):
            today_str = datetime.now().strftime("%d-%b-%Y")
            try:
                with open(FIXED_READING_FILE, "w") as f:
                    json.dump({"reading": prev_r, "date": today_str}, f)
                st.session_state.prev_reading = prev_r
                st.session_state.prev_date = today_str
                st.success("✅ Saved for next 30 days!")
                st.rerun()
            except Exception as e:
                st.error("Failed to save fixed reading.")
    
    st.markdown("##### 📸 Auto-Read (OCR)")
    uploaded_file = st.file_uploader("Upload or Take a Picture of the Meter", type=['jpg', 'jpeg', 'png'])
    if uploaded_file is not None:
        try:
            img = Image.open(uploaded_file)
            st.image(img, caption="Meter Preview", width=300)
            if OCR_AVAILABLE:
                text = pytesseract.image_to_string(img, config=r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789')
                numbers = re.findall(r'\d+', text)
                if numbers:
                    largest_num = int(max(numbers, key=len))
                    st.session_state.ocr_reading = largest_num
                    st.success(f"✅ OCR Extracted Number: **{largest_num}**")
                else:
                    st.warning("⚠ No number detected. Please enter manually.")
            else:
                st.info("Tesseract OCR is not installed/configured. Use manual entry.")
        except Exception as e:
            st.error("Error reading image.")
            
    c_curr_col, c_mon_col = st.columns(2)
    with c_curr_col:
        curr_r = st.number_input("Current Reading (New):", value=int(st.session_state.ocr_reading))
    with c_mon_col:
        month_str = st.text_input("Billing Month / Year:", value=datetime.now().strftime("%b-%Y"))
        
    add_chg = st.number_input("Additional Chg (Rs):", value=float(ADDITIONAL_CHG))
    
    st.markdown("<br>", unsafe_allow_html=True)
    if st.button("⚡ CALCULATE BILL", type="primary", use_container_width=True):
        if curr_r < prev_r:
            st.error("❌ Current reading cannot be less than previous.")
        else:
            units = curr_r - prev_r
            energy, fixed, additional, duty, salestax, kmc, total = calc_ke_bill(units, add_chg)
            st.session_state.computed = {
                "prev": prev_r, "curr": curr_r, "units": units,
                "energy": energy, "fixed": fixed, "additional": additional,
                "duty": duty, "salestax": salestax, "kmc": kmc, "total": total,
                "month": month_str, "cid": cid_input, "name": cname, "prop": cprop
            }

with col2:
    st.markdown("### 📊 Bill Summary")
    if st.session_state.computed:
        comp = st.session_state.computed
        
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-title">💡 UNITS CONSUMED</div>
            <div class="metric-value units">{comp['units']} kWh</div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-title">⚡ ENERGY CHARGES</div>
            <div class="metric-value">Rs. {comp['energy']:,.2f}</div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"""
        <div class="metric-card">
            <div class="metric-title">🧾 TAXES & OTHERS (Fixed, Duty, Tax, KMC, Add)</div>
            <div class="metric-value">Rs. {(comp['fixed'] + comp['duty'] + comp['salestax'] + comp['kmc'] + comp['additional']):,.2f}</div>
        </div>
        """, unsafe_allow_html=True)
        
        st.markdown(f"""
        <div class="metric-card" style="border-left-color: #dc3545; background-color: #fff5f5;">
            <div class="metric-title" style="color: #dc3545;">🚨 TOTAL PAYABLE BILL</div>
            <div class="metric-value total">Rs. {comp['total']:,.2f}</div>
        </div>
        """, unsafe_allow_html=True)

        st.markdown("---")
        st.markdown("### 📤 Actions")
        
        # Save to DB and Share functionality
        if st.button("💾 Save to Excel & Generate Docx", use_container_width=True):
            try:
                if not os.path.exists(DB_FILE):
                    wb = Workbook()
                    ws = wb.active
                    ws.title = "Electricity"
                    ws.append(["Rec_ID", "Cust_ID", "Customer", "Property", "Month",
                               "Prev_Read", "Curr_Read", "Units", "Energy_Charge",
                               "Fixed", "Additional", "Duty", "SalesTax", "KMC",
                               "Total_Bill", "KE_Acc", "Meter_No", "Date_Issued"])
                    wb.save(DB_FILE)

                wb = load_workbook(DB_FILE)
                if "Electricity" not in wb.sheetnames:
                    ws_el = wb.create_sheet("Electricity")
                    ws_el.append(["Rec_ID", "Cust_ID", "Customer", "Property", "Month",
                                  "Prev_Read", "Curr_Read", "Units", "Energy_Charge",
                                  "Fixed", "Additional", "Duty", "SalesTax", "KMC",
                                  "Total_Bill", "KE_Acc", "Meter_No", "Date_Issued"])
                else:
                    ws_el = wb["Electricity"]
                    
                ws_el.append([
                    ws_el.max_row, comp["cid"], comp["name"], comp["prop"], comp["month"],
                    comp["prev"], comp["curr"], comp["units"],
                    comp["energy"], comp["fixed"], comp["additional"],
                    comp["duty"], comp["salestax"], comp["kmc"],
                    comp["total"], "0400013715589", "SEA78555",
                    datetime.now().strftime("%d-%b-%Y %I:%M %p")
                ])
                wb.save(DB_FILE)
                st.success("✅ Saved to Excel Database globally!")
                
                # Generate Document
                now_str = datetime.now().strftime("%d-%b-%Y %I:%M %p")
                plain_txt = (
                    f"\\n{'='*55}\\n"
                    f"         {COMPANY_NAME}\\n"
                    f"          ELECTRICITY BILL RECEIPT\\n"
                    f"{'='*55}\\n"
                    f"  Generated On: {now_str}\\n"
                    f"  Month       : {comp['month']}\\n"
                    f"  Customer ID : {comp['cid']}\\n"
                    f"  Name        : {comp['name']}\\n"
                    f"  Property    : {comp['prop']}\\n"
                    f"{'-'*55}\\n"
                    f"  METER READINGS\\n"
                    f"  Previous Reading : {comp['prev']} kWh\\n"
                    f"  Current Reading  : {comp['curr']} kWh\\n"
                    f"  Units Consumed   : {comp['units']} kWh\\n"
                    f"{'-'*55}\\n"
                    f"  CHARGES\\n"
                    f"  Energy           : Rs. {comp['energy']:,.2f}\\n"
                    f"  Fixed Charge     : Rs. {comp['fixed']:,.2f}\\n"
                    f"  Additional       : Rs. {comp['additional']:,.2f}\\n"
                    f"  Taxes & Duty     : Rs. {(comp['duty'] + comp['salestax']):,.2f}\\n"
                    f"  KMC Charges      : Rs. {comp['kmc']:,.2f}\\n"
                    f"{'-'*55}\\n"
                    f"  TOTAL BILL       : Rs. {comp['total']:,.2f}\\n"
                    f"{'='*55}\\n"
                )
                
                doc = Document()
                doc.add_heading(f"{COMPANY_NAME} — Electricity Bill", 0)
                doc.add_paragraph(plain_txt.replace('\\n', '\n'))
                fname = os.path.join(_BASE_DIR, f"ElecBill_{sanitize_filename(comp['name'])}_{sanitize_filename(comp['month'])}.docx")
                doc.save(fname)
                st.success(f"✅ Downloadable Receipt Generated!")
                
                with open(fname, "rb") as file:
                    st.download_button(
                        label="📥 Download Docx Receipt",
                        data=file,
                        file_name=os.path.basename(fname),
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )
                
            except Exception as e:
                st.error(f"Error saving: {e}")

        st.markdown("---")
        wa_phone = st.text_input("📱 WhatsApp Number (Auto-Fill from DB):", value=st.session_state.customer_data.get('phone', ''))
        
        target_phone = wa_phone.replace(" ", "")
        if target_phone:
            if not target_phone.startswith("+"):
                target_phone = "+92" + target_phone.lstrip("0") if len(target_phone) >= 10 else "+" + target_phone
            
            # Using actual newlines for URL Encoding
            wa_msg = (
                f"⚡ *ELECTRICITY BILL — {COMPANY_NAME}*\n"
                f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                f"📅 *Month:* {comp['month']}\n"
                f"👤 *Customer:* {comp['name']}\n"
                f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                f"  Previous  : {comp['prev']} kWh\n"
                f"  Current   : {comp['curr']} kWh\n"
                f"  *Units Used : {comp['units']} kWh*\n"
                f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                f"💵 *TOTAL BILL : Rs. {comp['total']:,.2f}*\n"
                f"━━━━━━━━━━━━━━━━━━━━━━━━━━━━━\n"
                f"⌚ Time: {datetime.now().strftime('%d-%b-%Y %I:%M %p')}\n"
                f"⚠️ Please pay by the due date."
            )
            
            import urllib.parse
            encoded_msg = urllib.parse.quote(wa_msg)
            # wa.me requires phone number without the '+' sign
            clean_phone_for_url = target_phone.replace("+", "")
            wa_url = f"https://wa.me/{clean_phone_for_url}?text={encoded_msg}"
            
            st.markdown(
                f'<a href="{wa_url}" target="_blank" style="'
                f'display: block; text-align: center; background-color: #25D366; '
                f'color: white; padding: 10px; border-radius: 8px; font-weight: bold; '
                f'text-decoration: none; margin-top: 10px;">'
                f'📱 Open WhatsApp on mobile to send now!</a>', 
                unsafe_allow_html=True
            )
    else:
        st.info("👈 Enter details on the left and click **Calculate** to see the summary.")
